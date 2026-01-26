import streamlit as st
import pandas as pd
import random
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from io import BytesIO
from datetime import datetime

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Entrenador Pro Científico", layout="wide")

# --- GESTIÓN DE ESTADO ---
if 'reset_counter' not in st.session_state:
    st.session_state.reset_counter = 0

def get_key(base_name):
    return f"{base_name}_{st.session_state.reset_counter}"

# --- DATOS TEÓRICOS DE LOS OBJETIVOS ---
INFO_OBJETIVOS = {
    "Fuerza Máxima": """1️⃣ FUERZA MÁXIMA
🎯 Objetivo
Aumentar la capacidad máxima de producción de fuerza (adaptación neural).

🏋️‍♂️ Trabajo de fuerza
Intensidad: 85–100 % RM
Repeticiones: 1–5
Series: 4–6
Descanso: 3–6 min
Ejercicios: multiarticulares (sentadilla, peso muerto, press banca, press militar)

❤️ Trabajo cardiovascular
Tipo: aeróbico extensivo
Intensidad: 60–70 % FCmáx
Duración: 15–25 min
Frecuencia: 1–2 días/semana
Objetivo: recuperación, no interferir con la fuerza""",

    "Hipertrofia Muscular": """2️⃣ HIPERTROFIA MUSCULAR
🎯 Objetivo
Aumentar el tamaño muscular (hipertrofia miofibrilar y sarcoplasmática).

🏋️‍♂️ Trabajo de fuerza
Intensidad: 65–85 % RM
Repeticiones: 6–12
Series: 3–6
Descanso: 60–120 s
RIR: 0–2 (cerca del fallo)

❤️ Trabajo cardiovascular
Tipo: aeróbico moderado
Intensidad: 65–75 % FCmáx
Duración: 20–30 min
Frecuencia: 2–3 días/semana
Objetivo: salud cardiovascular sin comprometer ganancias musculares""",

    "Definición Muscular": """3️⃣ DEFINICIÓN MUSCULAR
🎯 Objetivo
Mantener masa muscular + reducir grasa corporal.

🏋️‍♂️ Trabajo de fuerza
Intensidad: 60–75 % RM
Repeticiones: 10–15
Series: 3–5
Descanso: 30–60 s
Métodos: superseries, circuitos, alta densidad

❤️ Trabajo cardiovascular
Tipo: HIIT + aeróbico
HIIT: 85–95 % FCmáx | 10–20 min | 1–2 días/sem
Aeróbico: 65–75 % FCmáx | 30–45 min | 2–3 días/sem""",

    "Resistencia Muscular": """4️⃣ RESISTENCIA MUSCULAR
🎯 Objetivo
Mejorar la capacidad de sostener esfuerzos prolongados.

🏋️‍♂️ Trabajo de fuerza
Intensidad: 30–60 % RM
Repeticiones: 15–30+
Series: 2–4
Descanso: 15–45 s
Formato: circuitos o estaciones

❤️ Trabajo cardiovascular
Tipo: aeróbico extensivo
Intensidad: 65–80 % FCmáx
Duración: 30–60 min
Frecuencia: 3–5 días/semana
Objetivo: base aeróbica y resistencia general""",

    "Mantenimiento Muscular": """5️⃣ MANTENIMIENTO MUSCULAR
🎯 Objetivo
Conservar masa muscular, fuerza y salud con bajo volumen.

🏋️‍♂️ Trabajo de fuerza
Intensidad: 60–75 % RM
Repeticiones: 8–12
Series: 2–3
Descanso: 60–90 s
Frecuencia: 2–3 días/semana

❤️ Trabajo cardiovascular
Tipo: aeróbico saludable
Intensidad: 60–75 % FCmáx
Duración: 20–40 min
Frecuencia: 2–4 días/semana""",

    "Rehabilitación Muscular y Articular": """5️⃣ REHABILITACIÓN MUSCULAR Y ARTICULAR
🔁 Progresión recomendada (por fases)

🟢 Fase 1 – Readaptación
20–30 % RM
Isométricos + movilidad
Cardio muy suave

🟡 Fase 2 – Reacondicionamiento
30–50 % RM
Concéntrico + excéntrico lento
Propiocepción dinámica

🔵 Fase 3 – Transición al entrenamiento
50–60 % RM
Patrones básicos
Integración progresiva con mantenimiento muscular""",

    "Programa de Pérdida de Peso": """🔥 PROGRAMA DE PÉRDIDA DE PESO

🎯 Objetivo
Reducir grasa corporal
Mantener o minimizar la pérdida de masa muscular
Aumentar el gasto energético total
Mejorar la salud metabólica y cardiovascular
Crear hábitos de actividad física sostenibles

🏋️‍♂️ Fuerza (entrenamiento principal)
🔹 Intensidad
50–70 % de 1RM
🔹 Repeticiones
12–20 repeticiones
🔹 Series
3–4 series
🔹 Descanso
20–45 segundos

🔹 Organización del trabajo
Circuitos
Superseries
Ejercicios multiarticulares prioritarios
Ritmo continuo, intensidad alta

Objetivo de la fuerza
Mantener masa muscular
Aumentar gasto calórico
Mejorar tono muscular

❤️ Entrenamiento cardiovascular
🔹 Aeróbico continuo
Intensidad: 60–75 % FCmáx
Duración: 30–60 min
Frecuencia: 3–5 días/semana
Ejemplos: caminar rápido, bici, elíptica, natación

🔹 HIIT (opcional)
Intensidad: 85–95 % FCmáx
Duración: 10–20 min
Frecuencia: 1–2 días/semana
Formato: intervalos cortos de alta intensidad + recuperación activa

🧠 Consejos clave
La fuerza es imprescindible para no perder músculo
No bajar de 50 % RM de forma sistemática
Mantener déficit calórico moderado
Priorizar adherencia y progresión
Dormir y recuperarse adecuadamente
Aumentar el NEAT- Non Exercice Activity Thermogenesis (pasos diarios, vida activa)
Revaluar cargas cada 4–6 semanas

⚠️ Errores comunes
Solo cardio y nada de fuerza
Usar cargas muy ligeras durante meses
Descansos excesivos
Déficits calóricos extremos"""
}

# --- FUNCIONES AUXILIARES WORD ---
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_cell_bg_color(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_header_cell(cell, text, width_inches=None):
    cell.text = text
    if width_inches:
        cell.width = Inches(width_inches)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg_color(cell, "2E4053")

def set_row_cant_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)

def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True

# --- BUSCADOR ---
def encontrar_imagen_recursiva(nombre_objetivo):
    if not nombre_objetivo or pd.isna(nombre_objetivo):
        return None, "Celda Vacía"
    nombre_limpio = str(nombre_objetivo).strip().lower()
    nombre_base_excel = os.path.splitext(nombre_limpio)[0]
    for root, dirs, files in os.walk("."):
        for filename in files:
            if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
                filename_base = os.path.splitext(filename)[0].lower()
                if filename.lower() == nombre_limpio:
                    return os.path.join(root, filename), "Exacta"
                if filename_base == nombre_base_excel:
                    return os.path.join(root, filename), "Por Nombre"
    return None, f"No encontrado"

# --- CARGAR EXCEL ---
@st.cache_data
def cargar_ejercicios():
    try:
        if os.path.exists("DB_EJERCICIOS.xlsx"):
            df = pd.read_excel("DB_EJERCICIOS.xlsx")
            df.columns = df.columns.str.strip().str.lower()
            if 'nombre' not in df.columns:
                if 'ejercicio' in df.columns: df.rename(columns={'ejercicio': 'nombre'}, inplace=True)
            for col in ['tipo', 'imagen', 'desc', 'agonistas', 'sinergistas', 'estabilizadores']:
                if col not in df.columns: df[col] = ""
            
            df['tipo'] = df['tipo'].astype(str).str.replace('Olimpica', 'Olímpica', regex=False)
            df['tipo'] = df['tipo'].str.replace('olimpica', 'Olímpica', regex=False, case=False)
            df['tipo'] = df['tipo'].str.replace('Rehabilitacion', 'Rehabilitación', regex=False)
            df['tipo'] = df['tipo'].str.replace('Rotualiana', 'Rotuliana', regex=False)
            df['tipo'] = df['tipo'].str.strip()
            df = df.fillna("")
            return df.to_dict('records')
        else:
            return None
    except Exception as e:
        return f"Error: {str(e)}"

DB_EJERCICIOS = cargar_ejercicios()

# --- GENERADOR WORD ---
def generar_word_final(rutina_df, lista_estiramientos, objetivo, alumno, titulo_material, intensidad_str, cardio_tipo, cardio_tiempo, series_str, incluir_analisis_muscular):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    # --- PIE DE PÁGINA PROFESIONAL (OPTIMIZADO) ---
    footer = section.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    
    # Tabla Footer 1x2 - Ajustada para no ocupar mucho espacio vertical
    ft = footer.add_table(rows=1, cols=2, width=Inches(10.8))
    ft.autofit = False
    ft.columns[0].width = Inches(4.0) 
    ft.columns[1].width = Inches(6.8) 
    
    # Celda Izq: Marca de agua
    path_watermark, _ = encontrar_imagen_recursiva("logo_firma")
    if not path_watermark:
        path_watermark, _ = encontrar_imagen_recursiva("watermark")

    if path_watermark:
        c_logo = ft.cell(0,0)
        c_logo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_logo = c_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_logo = p_logo.add_run()
        # --- CAMBIO CRUCIAL: TAMAÑO REDUCIDO ---
        r_logo.add_picture(path_watermark, height=Inches(0.55)) 
    
    # Celda Der: Número de Página
    c_txt = ft.cell(0,1)
    c_txt.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER # Alineado al centro vertical
    p_txt = c_txt.paragraphs[0]
    p_txt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run_pag_label = p_txt.add_run("Página ")
    run_pag_label.font.size = Pt(10)
    
    run_num = p_txt.add_run()
    run_num.font.size = Pt(10)
    add_page_number(run_num)

    # --- CUERPO DEL DOCUMENTO ---
    
    # PÁGINA 1
    head_tbl = doc.add_table(rows=1, cols=2)
    head_tbl.autofit = False
    head_tbl.columns[0].width = Inches(9.8) 
    head_tbl.columns[1].width = Inches(1.0)
    
    c1 = head_tbl.cell(0,0)
    p = c1.paragraphs[0]
    r1 = p.add_run(f"PROGRAMA DE ENTRENAMIENTO DE: {titulo_material.upper()}\n")
    r1.font.bold = True
    r1.font.size = Pt(12) 
    r1.font.color.rgb = RGBColor(41, 128, 185)
    
    nombre_mostrar = alumno if alumno.strip() else "ALUMNO"
    
    font_size_meta = Pt(10) 
    r_obj_label = p.add_run("OBJETIVO: ")
    r_obj_label.font.bold = True
    r_obj_label.font.size = font_size_meta
    p.add_run(f"{objetivo}").font.size = font_size_meta
    p.add_run("   |   ").font.size = font_size_meta 
    
    r_int_label = p.add_run("INTENSIDAD DE TRABAJO: ")
    r_int_label.font.bold = True
    r_int_label.font.size = font_size_meta
    p.add_run(f"({intensidad_str})").font.size = font_size_meta
    p.add_run("   |   ").font.size = font_size_meta 
    
    r_alu_label = p.add_run("ALUMNO/A: ")
    r_alu_label.font.bold = True
    r_alu_label.font.size = font_size_meta
    p.add_run(f"{nombre_mostrar.upper()}").font.size = font_size_meta

    c2 = head_tbl.cell(0,1)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run(f"FECHA:\n{datetime.now().strftime('%d/%m/%Y')}").bold = True
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    run_sub = p_sub.add_run("Situación de Aprendizaje: Trabajo en Salas de Musculación 1º de Bachillerato IES Lucía de Medrano")
    run_sub.font.bold = True
    run_sub.font.name = 'Cambria'
    run_sub.font.size = Pt(16)    
    rPr = run_sub._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Cambria')
    rFonts.set(qn('w:hAnsi'), 'Cambria')
    rPr.append(rFonts)

    doc.add_paragraph("_" * 95)

    h1 = doc.add_heading(level=1)
    titulo_seccion_1 = '1. Guía Visual de Ejercicios con Análisis Muscular' if incluir_analisis_muscular else '1. Guía Visual de Ejercicios'
    run_h1 = h1.add_run(titulo_seccion_1)
    run_h1.font.size = Pt(18)
    run_h1.font.color.rgb = RGBColor(44, 62, 80)

    cardio_table = doc.add_table(rows=1, cols=2)
    cardio_table.style = 'Table Grid'
    c_warm = cardio_table.cell(0,0)
    p_warm = c_warm.paragraphs[0]
    p_warm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_w = p_warm.add_run("A) Calentamiento de 5 minutos de Duración")
    run_w.font.bold = True
    run_w.font.size = Pt(10)
    set_cell_bg_color(c_warm, "EAEDED") 
    c_card = cardio_table.cell(0,1)
    p_card = c_card.paragraphs[0]
    p_card.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = p_card.add_run(f"B) Cardio: {cardio_tipo} -> {cardio_tiempo}")
    run_c.font.bold = True
    run_c.font.size = Pt(10) 
    set_cell_bg_color(c_card, "EAEDED") 
    
    doc.add_paragraph("")

    num_ej = len(rutina_df)
    cols_visual = 4
    rows_visual = (num_ej + cols_visual - 1) // cols_visual
    vis_table = doc.add_table(rows=rows_visual, cols=cols_visual)
    vis_table.style = 'Table Grid'
    
    TR_HEIGHT_TWIPS = 3800 if incluir_analisis_muscular else 2800 
    
    for row in vis_table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(TR_HEIGHT_TWIPS))
        trHeight.set(qn('w:hRule'), "atLeast")
        trPr.append(trHeight)
        set_row_cant_split(row)

    for i, row_data in enumerate(rutina_df.to_dict('records')):
        r = i // cols_visual
        c = i % cols_visual
        cell = vis_table.cell(r, c)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        ruta_img, msg = encontrar_imagen_recursiva(row_data['Imagen'])
        if ruta_img:
            try:
                run = p.add_run()
                run.add_picture(ruta_img, width=Inches(2.4), height=Inches(1.55))
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
            except:
                p.add_run(f"[Error]\n")
        else:
            p.add_run(f"\n[FOTO NO DISPONIBLE]\n")
        
        run_nom = p.add_run("\n" + row_data['Ejercicio'])
        run_nom.font.bold = True
        run_nom.font.size = Pt(10)
        
        if incluir_analisis_muscular:
            p.add_run("\n" + "_"*25 + "\n").font.size = Pt(6)
            
            p.add_run("Músculos Agonistas:\n").font.bold = True
            p.add_run(f"{str(row_data.get('agonistas', ''))}\n").font.size = Pt(8)
            
            p.add_run("Músculos Sinergistas:\n").font.bold = True
            p.add_run(f"{str(row_data.get('sinergistas', ''))}\n").font.size = Pt(8)

            p.add_run("Músculos Estabilizadores:\n").font.bold = True
            p.add_run(f"{str(row_data.get('estabilizadores', ''))}").font.size = Pt(8)

    doc.add_page_break()

    # PÁGINA 2
    h2 = doc.add_heading(level=1)
    run_h2 = h2.add_run('2. Rutina Detallada')
    run_h2.font.size = Pt(18)
    run_h2.font.color.rgb = RGBColor(44, 62, 80)

    tech_table = doc.add_table(rows=1, cols=6)
    tech_table.style = 'Table Grid'
    tech_table.autofit = False 
    widths = [0.7, 3.5, 1.5, 1.0, 1.5, 2.4] 
    headers = ["Orden", "Ejercicio", "Series x Reps", "Carga", "Descanso", "Notas"]
    row_hdr = tech_table.rows[0]
    set_row_cant_split(row_hdr)
    for i, h in enumerate(headers):
        style_header_cell(row_hdr.cells[i], h, widths[i])
        
    for idx, row_data in rutina_df.iterrows():
        row_cells = tech_table.add_row().cells
        set_row_cant_split(tech_table.rows[-1])
        for i in range(6):
            row_cells[i].width = Inches(widths[i])
        row_cells[0].text = str(idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = row_data['Ejercicio']
        row_cells[2].text = f"{series_str} x {row_data['Reps']}"
        row_cells[3].text = f"{row_data['Peso']} kg"
        row_cells[4].text = row_data['Descanso']
        row_cells[5].text = f"Int: {row_data['Intensidad_Real']}" 

    doc.add_paragraph("\n")

    if lista_estiramientos:
        h3 = doc.add_heading(level=1)
        run_h3 = h3.add_run('3. Ejercicios de Estiramientos')
        run_h3.font.size = Pt(18)
        run_h3.font.color.rgb = RGBColor(44, 62, 80)
        h3.paragraph_format.keep_with_next = True

        num_est = len(lista_estiramientos)
        cols_est = 4
        rows_est = (num_est + cols_est - 1) // cols_est
        est_table = doc.add_table(rows=rows_est, cols=cols_est)
        est_table.style = 'Table Grid'
        for row in est_table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(2600))
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
            set_row_cant_split(row)

        for i, item_est in enumerate(lista_estiramientos):
            r = i // cols_est
            c = i % cols_est
            cell = est_table.cell(r, c)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ruta_img, msg = encontrar_imagen_recursiva(item_est['imagen'])
            if ruta_img:
                try:
                    run = p.add_run()
                    run.add_picture(ruta_img, width=Inches(2.2), height=Inches(1.4))
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                except:
                    p.add_run(f"[Error]\n")
            else:
                p.add_run(f"\n[FOTO NO DISPONIBLE]\n")
            run_nom = p.add_run("\n" + item_est['nombre'])
            run_nom.font.bold = True
            run_nom.font.size = Pt(9)
        doc.add_paragraph("\n")

    # ================= SECCIÓN 4: BORG (BLOQUE INDIVISIBLE) =================
    h4 = doc.add_heading(level=1)
    run_h4 = h4.add_run('4. Percepción del Esfuerzo (RPE) - Escala de Borg')
    run_h4.font.size = Pt(18)
    run_h4.font.color.rgb = RGBColor(44, 62, 80)
    set_keep_with_next(h4)

    borg_table = doc.add_table(rows=3, cols=5)
    borg_table.style = 'Table Grid'
    borg_table.autofit = True
    
    for row in borg_table.rows:
        set_row_cant_split(row)

    borg_data = [
        {"val": "6-8", "txt": "Muy Ligero", "icon": "🙂", "color": "A9DFBF"},
        {"val": "9-11", "txt": "Ligero", "icon": "😌", "color": "D4EFDF"},
        {"val": "12-14", "txt": "Algo Duro", "icon": "😐", "color": "F9E79F"},
        {"val": "15-17", "txt": "Duro", "icon": "😓", "color": "F5CBA7"},
        {"val": "18-20", "txt": "Máximo", "icon": "🥵", "color": "E6B0AA"}
    ]
    
    row_icons = borg_table.rows[0]
    for i, data in enumerate(borg_data):
        c = row_icons.cells[i]
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_icon = p.add_run(f"{data['icon']}\n")
        run_icon.font.size = Pt(26) 
        run_val = p.add_run(f"{data['val']}")
        run_val.font.size = Pt(14)
        set_cell_bg_color(c, data['color'])

    row_text = borg_table.rows[1]
    for i, data in enumerate(borg_data):
        c = row_text.cells[i]
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(data['txt']).font.bold = True
        set_cell_bg_color(c, data['color'])

    row_check = borg_table.rows[2]
    tr = row_check._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "600")
    trPr.append(trHeight)
    for i, data in enumerate(borg_data):
        c = row_check.cells[i]
        set_cell_bg_color(c, data['color'])

    p_note = doc.add_paragraph("Marca con una X la sensación global al terminar el entrenamiento.")
    p_note.style = "Caption"

    doc.add_paragraph("\n")

    # --- SECCIÓN 5: MARCO TEÓRICO ---
    h5 = doc.add_heading(level=1)
    run_h5 = h5.add_run(f"5. {objetivo.upper()}")
    run_h5.font.size = Pt(18)
    run_h5.font.color.rgb = RGBColor(44, 62, 80)
    
    raw_text = INFO_OBJETIVOS.get(objetivo, "Información no disponible.")
    clean_lines = raw_text.split('\n')[1:] 
    
    emojis_clave = ['🎯', '🏋️‍♂️', '❤️', '🔁', '🟢', '🟡', '🔵', '🔥', '🔹', '🧠', '⚠️']
    
    for line in clean_lines:
        if not line.strip(): 
            continue
