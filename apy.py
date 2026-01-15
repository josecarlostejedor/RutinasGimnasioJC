import streamlit as st
import pandas as pd
import random
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from io import BytesIO
from datetime import datetime

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Generador Rutinas con Imágenes", layout="wide")

# --- CARGAR BASE DE DATOS ---
@st.cache_data
def cargar_ejercicios():
    try:
        if os.path.exists("DB_EJERCICIOS.xlsx"):
            df = pd.read_excel("DB_EJERCICIOS.xlsx")
            # Rellenar valores nulos en la columna imagen por si acaso
            if 'imagen' not in df.columns:
                df['imagen'] = ""
            df['imagen'] = df['imagen'].fillna("")
            return df.to_dict('records')
        else:
            return None
    except Exception as e:
        return str(e)

DB_EJERCICIOS = cargar_ejercicios()

# --- LÓGICA DE PARÁMETROS ---
def obtener_parametros(objetivo):
    # Ajusta estos valores según tu criterio profesional
    if objetivo == "Hipertrofia":
        return {"reps": "6-12", "int_min": 0.60, "int_max": 0.85, "descanso": "1-3 min"}
    elif objetivo == "Fuerza Máxima":
        return {"reps": "1-5", "int_min": 0.85, "int_max": 0.95, "descanso": "3-5 min"}
    elif objetivo == "Resistencia":
        return {"reps": "15-20", "int_min": 0.40, "int_max": 0.60, "descanso": "30-60 s"}

# --- GENERADOR DE WORD ESTILO PDF (HORIZONTAL + IMÁGENES) ---
def generar_word_con_imagenes(rutina_df, objetivo, alumno):
    doc = Document()
    
    # 1. Configurar página Horizontal
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # 2. Encabezado
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    header_table.columns[0].width = Inches(4)
    header_table.columns[2].width = Inches(2.5)
    
    header_table.cell(0,0).text = "PLAN DE ENTRENAMIENTO PERSONALIZADO"
    header_table.cell(0,1).text = f"OBJETIVO: {objetivo.upper()}"
    c3 = header_table.cell(0,2)
    c3.text = f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\nIES / CLUB DEPORTIVO"
    c3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph(f"Alumno/a: {alumno}")
    doc.add_paragraph("_" * 115)

    # 3. Estructura Principal (2 Columnas: Visual vs Registro)
    main_table = doc.add_table(rows=1, cols=2)
    main_table.autofit = False
    main_table.columns[0].width = Inches(5.5) # Panel Imágenes
    main_table.columns[1].width = Inches(5.0) # Panel Tabla

    # --- COLUMNA IZQUIERDA: IMÁGENES ---
    left_cell = main_table.cell(0,0)
    
    # Tabla anidada para grid de imágenes (2 columnas internas)
    filas_necesarias = (len(rutina_df) + 1) // 2
    visual_table = left_cell.add_table(rows=filas_necesarias, cols=2)
    visual_table.style = 'Table Grid'
    
    for idx, row_data in rutina_df.iterrows():
        r = idx // 2
        c = idx % 2
        cell = visual_table.cell(r, c)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # INSERTAR IMAGEN
        nombre_archivo = str(row_data['Imagen_Archivo']).strip()
        ruta_imagen = os.path.join("images", nombre_archivo)
        
        # Verificar si existe la imagen en la carpeta
        if nombre_archivo and os.path.exists(ruta_imagen):
            try:
                run = p.add_run()
                run.add_picture(ruta_imagen, width=Inches(1.8)) # Ajustar tamaño
                p.add_run("\n")
            except Exception:
                p.add_run("[Error Imagen]\n")
        else:
            # Si no hay imagen, poner un placeholder
            p.add_run("\n[FOTO]\n")

        # Nombre del ejercicio bajo la foto
        run_text = p.add_run(row_data['Ejercicio'])
        run_text.font.bold = True
        run_text.font.size = Pt(9)

    # --- COLUMNA DERECHA: TABLA DE DATOS ---
    right_cell = main_table.cell(0,1)
    
    reg_table = right_cell.add_table(rows=1, cols=3)
    reg_table.style = 'Table Grid'
    
    # Cabeceras
    hdr = reg_table.rows[0].cells
    hdr[0].text = "Ejercicio"
    hdr[1].text = "Series/Reps"
    hdr[2].text = "Kg"
    
    for idx, row_data in rutina_df.iterrows():
        # Fila datos
        row = reg_table.add_row().cells
        row[0].text = f"{idx + 1}. {row_data['Ejercicio']}"
        row[1].text = f"{row_data['Series']} x {row_data['Reps']}"
        row[2].text = str(row_data['Peso'])
        
        # Fila descanso
        row_d = reg_table.add_row().cells
        merged = row_d[0].merge(row_d[2])
        merged.text = f"Descanso: {row_data['Descanso']} | Notas: _________________"
        merged.paragraphs[0].runs[0].font.size = Pt(8)

    # 4. Footer (Borg)
    doc.add_paragraph("\n")
    footer = doc.add_table(rows=1, cols=2)
    footer.style = 'Table Grid'
    c_borg = footer.cell(0,1)
    c_borg.text = "Escala de Borg: 6-7 (Muy ligero) | 12-13 (Algo duro) | 18-20 (Agotamiento)"
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFAZ STREAMLIT ---
st.title("🏋️ Generador de Rutinas (Estilo Visual)")

if DB_EJERCICIOS is None:
    st.error("⚠️ Error: No se encuentra 'DB_EJERCICIOS.xlsx'. Súbelo al GitHub.")
    st.stop()
elif isinstance(DB_EJERCICIOS, str):
    st.error(f"Error leyendo Excel: {DB_EJERCICIOS}")
    st.stop()

# Sidebar
st.sidebar.header("Configuración")
alumno = st.sidebar.text_input("Nombre Alumno:", "Atleta")
objetivo = st.sidebar.selectbox("Objetivo", ["Hipertrofia", "Fuerza Máxima", "Resistencia"])

# Filtrar número de ejercicios
num_total = len(DB_EJERCICIOS)
num_ej = st.sidebar.slider("Nº Ejercicios", 1, min(10, num_total), min(6, num_total))

lista_nombres = [e['nombre'] for e in DB_EJERCICIOS]
seleccion = st.multiselect("Seleccionar ejercicios específicos:", lista_nombres, max_selections=num_ej)

# Lógica autocompletar
ejercicios_finales_data = []
nombres_seleccionados = seleccion.copy()

# Si faltan ejercicios para llegar al número del slider, añadir aleatorios
if len(nombres_seleccionados) < num_ej:
    restantes = [e for e in DB_EJERCICIOS if e['nombre'] not in nombres_seleccionados]
    faltan = num_ej - len(nombres_seleccionados)
    if faltan > 0:
        extra = random.sample(restantes, faltan)
        nombres_seleccionados.extend([e['nombre'] for e in extra])

# Recuperar datos completos de los nombres seleccionados
for nombre in nombres_seleccionados:
    item = next(x for x in DB_EJERCICIOS if x['nombre'] == nombre)
    ejercicios_finales_data.append(item)

# Inputs de RM
st.subheader("Cargas (RM)")
cols = st.columns(3)
inputs_rm = {}
for i, item in enumerate(ejercicios_finales_data):
    with cols[i % 3]:
        inputs_rm[item['nombre']] = st.number_input(f"RM {item['nombre']}", value=50, step=5)

if st.button("Generar PDF (Word)"):
    params = obtener_parametros(objetivo)
    rutina_export = []
    
    for item in ejercicios_finales_data:
        rm = inputs_rm[item['nombre']]
        # Cálculo simple de peso según % aleatorio del rango del objetivo
        intensidad = random.uniform(params['int_min'], params['int_max'])
        peso = round(rm * intensidad)
        
        rutina_export.append({
            "Ejercicio": item['nombre'],
            "Imagen_Archivo": item['imagen'], # Dato del Excel
            "Series": 4, # Simplificado o random
            "Reps": params['reps'],
            "Peso": peso,
            "Descanso": params['descanso']
        })
        
    df_export = pd.DataFrame(rutina_export)
    
    # Generar
    docx = generar_word_con_imagenes(df_export, objetivo, alumno)
    
    st.success("Rutina generada. Las imágenes del Excel se han insertado.")
    st.download_button(
        "📥 Descargar Word", 
        docx, 
        f"rutina_{alumno}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )